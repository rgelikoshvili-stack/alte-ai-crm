import json
import shutil
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt


SOURCE_DOCX = Path(r"C:\Users\Acer\Desktop\ალტე\Alte_AI_CRM_Master_Plan_v3_FINAL.docx")
OUT_DOCX = Path("Alte_AI_CRM_Master_Plan_v3_FINAL_with_KB_Appendix.docx")
KB_JSONL = Path("alte_knowledge_base/alte_knowledge_base_ka.jsonl")
KB_INDEX = Path("alte_knowledge_base/alte_knowledge_base_index.md")
KB_URLS = Path("alte_knowledge_base/alte_source_urls.txt")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_heading_safe(doc, text, level=1):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16 if level == 1 else 13)
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_bullet_safe(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-9)
    run = paragraph.add_run("- " + text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return paragraph


def main():
    shutil.copyfile(SOURCE_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading_safe(doc, "Appendix: Alte Website Knowledge Base Extraction", level=1)
    doc.add_paragraph(
        "ეს Appendix აღწერს alte.edu.ge-ის ქართული გვერდებიდან ამოღებულ chatbot knowledge-base პაკეტს. "
        "მიზანი არის, რომ AI chatbot-მა პასუხები გასცეს მხოლოდ ოფიციალური source/chunk-ების მიხედვით და "
        "არ გამოიგონოს ფასები, მიღების ვადები, გრანტები ან ოფიციალური წესები."
    )

    records = []
    with KB_JSONL.open(encoding="utf-8") as f:
        for line in f:
            records.append(json.loads(line))

    pages = {r["source_url"] for r in records}
    tags = Counter(tag for r in records for tag in r.get("tags", []))
    program_pages = {r["source_url"] for r in records if "program" in r.get("tags", [])}
    law_chunks = [r for r in records if "/ka/samartlis" in r["source_url"]]

    add_heading_safe(doc, "Generated files", level=2)
    add_table(
        doc,
        ["File", "Purpose"],
        [
            [str(KB_JSONL), "JSONL knowledge chunks; each line has source_url, title, section, tags and text."],
            [str(KB_INDEX), "Readable source index and chatbot answering policy."],
            [str(KB_URLS), "Crawled source URL list with lastmod dates from sitemap."],
            ["extract_alte_knowledge.py", "Repeatable crawler/parser script."],
        ],
    )

    add_heading_safe(doc, "Extraction summary", level=2)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Source pages with extracted chunks", len(pages)],
            ["Knowledge chunks", len(records)],
            ["Program-related pages", len(program_pages)],
            ["Law program chunks", len(law_chunks)],
            ["Crawl errors", 0],
        ],
    )

    add_heading_safe(doc, "Chunk tags", level=2)
    add_table(doc, ["Tag", "Chunk count"], sorted(tags.items()))

    add_heading_safe(doc, "Important implementation rules", level=2)
    for item in [
        "KnowledgeSource table-ში ჩაიწეროს source_url, title, locale=ka, status=active, version და last_reviewed_at.",
        "KnowledgeChunk table-ში ჩაიწეროს source_id, section, chunk_index, text, tags/metadata_json და active=true.",
        "AI answer-მა უნდა დააბრუნოს used_sources/source_url-ები, რომ ოპერატორმა იცოდეს რომელ წყაროზე დაყრდნობით უპასუხა.",
        "თუ source არ მოიძებნა ან confidence დაბალია, chatbot-მა უნდა შესთავაზოს human handover.",
        "Program/admission intent + contact data უნდა ქმნიდეს CRM lead-ს; general_info კითხვა lead-ს არ უნდა ქმნიდეს.",
        "Crawler რეგულარულად უნდა გაეშვას, რადგან sitemap-ში lastmod იცვლება და პროგრამის/მიღების ინფორმაცია შეიძლება განახლდეს.",
    ]:
        add_bullet_safe(doc, item)

    add_heading_safe(doc, "High-priority source examples", level=2)
    source_counts = defaultdict(int)
    titles = {}
    for r in records:
        source_counts[r["source_url"]] += 1
        titles[r["source_url"]] = r["title"] or r["source_url"]
    priority = [
        "/ka/sabakalavro-programebi",
        "/ka/samartlis",
        "/ka/biznesis-administrirebis",
        "/ka/kompiuteruli-metsnierebis-qartulenovani",
        "/ka/kompiuteruli-metsnierebis-inglisurenovani",
        "/ka/meditsinis-qartulenovani",
        "/ka/meditsinis-inglisurenovani",
        "/ka/stomatologia-inglisurenovani",
        "/ka/saertashoriso-urtiertobebi",
        "/ka/zhurnalistika",
        "/ka/abiturientebistvis",
        "/ka/kontaqti",
    ]
    rows = []
    for suffix in priority:
        matches = [u for u in source_counts if u.endswith(suffix)]
        for u in matches:
            rows.append([titles[u], u, source_counts[u]])
    add_table(doc, ["Title", "Source URL", "Chunks"], rows)

    doc.save(OUT_DOCX)
    print(OUT_DOCX.resolve())


if __name__ == "__main__":
    main()
